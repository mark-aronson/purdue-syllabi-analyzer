"""Analyze syllabi using the Claude API."""

import anthropic
import base64
import json
import re
import sys
from pathlib import Path

import docx
from dotenv import load_dotenv

load_dotenv()

from .models import SyllabusReview

PROMPT_PATH = Path(__file__).parent / "prompt.txt"
SYLLABI_DIR = Path(__file__).parents[2] / "syllabi"
RESULTS_DIR = Path(__file__).parents[2] / "data" / "results"
PROGRAMS_FILE = Path(__file__).parents[2] / "data" / "programs.json"


def load_system_prompt() -> str:
    return PROMPT_PATH.read_text(encoding="utf-8")


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


_DEPT_RE = re.compile(r"^([A-Z]+)\d")


def _extract_department(filename: str) -> str | None:
    """Extract the department code from a syllabus filename (e.g. 'POL10100_Spring2026_X.pdf' -> 'POL')."""
    m = _DEPT_RE.match(filename)
    return m.group(1) if m else None


def _extract_docx_text(docx_path: Path) -> str:
    """Extract plain text from a DOCX file."""
    doc = docx.Document(str(docx_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def _build_message_content(file_path: Path) -> list[dict]:
    """Build the Claude API message content blocks for a syllabus file."""
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        pdf_b64 = base64.standard_b64encode(file_path.read_bytes()).decode("utf-8")
        return [
            {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": pdf_b64,
                },
            },
            {
                "type": "text",
                "text": "Analyze this syllabus and extract the requested information as JSON.",
            },
        ]
    elif suffix == ".docx":
        text = _extract_docx_text(file_path)
        return [
            {
                "type": "text",
                "text": (
                    f"Below is the text content of a syllabus document ({file_path.name}):\n\n"
                    f"{text}\n\n"
                    "Analyze this syllabus and extract the requested information as JSON."
                ),
            },
        ]
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def analyze_syllabus(client: anthropic.Anthropic, file_path: Path) -> SyllabusReview:
    """Send a single syllabus (PDF or DOCX) to Claude for analysis and return a validated SyllabusReview."""
    message = client.messages.create(
        model="claude-sonnet-4-5-20250929",
        max_tokens=4096,
        system=load_system_prompt(),
        messages=[
            {
                "role": "user",
                "content": _build_message_content(file_path),
            }
        ],
    )

    text = message.content[0].text.strip()
    # Strip markdown code fences if present
    if text.startswith("```"):
        text = text.split("\n", 1)[1]  # remove opening ```json line
        text = text.rsplit("```", 1)[0]  # remove closing ```
    raw = json.loads(text)
    return SyllabusReview.model_validate(raw)


def _save_results(results: list[dict], output_path: Path) -> None:
    """Write the current results list to disk."""
    output_path.write_text(json.dumps(results, indent=2), encoding="utf-8")


def _analyze_files(client: anthropic.Anthropic, files: list[Path], output_path: Path) -> list[dict]:
    """Analyze a list of syllabi files, saving results incrementally to output_path."""
    # Load any existing results to support resuming
    if output_path.exists():
        results = json.loads(output_path.read_text(encoding="utf-8"))
        done = {r["_source_file"] for r in results}
    else:
        results = []
        done = set()

    for file_path in files:
        rel = file_path.name
        if rel in done:
            print(f"  Skipping {file_path.name} (already analyzed)")
            continue
        print(f"  Analyzing {file_path.name}...")
        try:
            review = analyze_syllabus(client, file_path)
            result = review.model_dump()
            result["_source_file"] = rel
            results.append(result)
        except Exception as e:
            print(f"  ERROR processing {file_path.name}: {e}")
            results.append({"_source_file": rel, "_error": str(e)})
        _save_results(results, output_path)

    return results


def analyze_department(client: anthropic.Anthropic, department: str) -> list[dict]:
    """Analyze all syllabi (PDF and DOCX) for a department in the flat syllabi directory."""
    files = sorted(
        f
        for f in SYLLABI_DIR.iterdir()
        if f.suffix.lower() in SUPPORTED_EXTENSIONS
        and _extract_department(f.name) == department
    )
    if not files:
        print(f"No supported files found for department {department}")
        return []

    RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    return _analyze_files(client, files, RESULTS_DIR / f"{department}.json")


def _normalize_course_code(course: str) -> str:
    """Normalize a course code to match the syllabi filename prefix format.

    Removes spaces, uppercases, and pads the course number to 5 digits with
    trailing zeros (e.g. "ABE 201" -> "ABE20100", "POL 10100" -> "POL10100").
    """
    code = course.replace(" ", "").upper()
    m = re.match(r"^([A-Z]+)(\d+)$", code)
    if not m:
        return code
    dept, num = m.group(1), m.group(2)
    return dept + num.ljust(5, "0")


def find_missing_syllabi(program_name: str) -> list[str] | None:
    """Find courses in a program that have no syllabi in the syllabi directory.

    Returns a sorted list of course codes (in their original format from
    programs.json) that have no matching files, or None if the program is not
    found.
    """
    if not PROGRAMS_FILE.exists():
        print(f"Programs file not found: {PROGRAMS_FILE}")
        return None

    programs = json.loads(PROGRAMS_FILE.read_text(encoding="utf-8"))
    if program_name not in programs:
        print(f"Program not found: {program_name}")
        print(f"Available programs: {', '.join(programs.keys())}")
        return None

    courses = programs[program_name]

    # Collect all course codes present in the syllabi directory
    syllabi_codes = {
        f.stem.split("_")[0].upper()
        for f in SYLLABI_DIR.iterdir()
        if f.suffix.lower() in SUPPORTED_EXTENSIONS
    }

    missing = [c for c in courses if _normalize_course_code(c) not in syllabi_codes]
    missing.sort()

    print(f"Program '{program_name}': {len(courses)} courses, "
          f"{len(courses) - len(missing)} found, {len(missing)} missing")
    if missing:
        for c in missing:
            print(f"  {c}")

    # Save results to JSON
    missing_dir = Path(__file__).parents[2] / "data" / "missing"
    missing_dir.mkdir(parents=True, exist_ok=True)
    slug = program_name.replace(" ", "_")
    output_path = missing_dir / f"{slug}.json"
    output = {
        "program": program_name,
        "total_courses": len(courses),
        "found": len(courses) - len(missing),
        "missing_count": len(missing),
        "missing_courses": missing,
    }
    output_path.write_text(json.dumps(output, indent=2), encoding="utf-8")
    print(f"Results saved to {output_path}")

    return missing


def analyze_program(program_name: str):
    """Analyze all syllabi for courses in a program defined in programs.json.

    Results are saved to per-department result files so the dashboard can display them.
    """
    if not PROGRAMS_FILE.exists():
        print(f"Programs file not found: {PROGRAMS_FILE}")
        return

    programs = json.loads(PROGRAMS_FILE.read_text(encoding="utf-8"))
    if program_name not in programs:
        print(f"Program not found: {program_name}")
        print(f"Available programs: {', '.join(programs.keys())}")
        return

    # Convert course codes like "POL 10100" to filename prefixes like "POL10100"
    course_prefixes = {c.replace(" ", "").upper() for c in programs[program_name]}

    # Find matching syllabi files
    all_files = sorted(
        f
        for f in SYLLABI_DIR.iterdir()
        if f.suffix.lower() in SUPPORTED_EXTENSIONS
    )

    # Match files whose name starts with a course prefix (before the first underscore)
    matching_files: list[Path] = []
    for f in all_files:
        # Extract course code from filename (e.g. "POL10100" from "POL10100_Spring2026_X.pdf")
        code = f.stem.split("_")[0].upper()
        if code in course_prefixes:
            matching_files.append(f)

    if not matching_files:
        print(f"No syllabi files found for program '{program_name}'")
        print(f"  Looking for files matching {len(course_prefixes)} course codes")
        return

    # Group files by department and analyze into per-department result files
    dept_files: dict[str, list[Path]] = {}
    for f in matching_files:
        dept = _extract_department(f.name)
        if dept:
            dept_files.setdefault(dept, []).append(f)

    client = anthropic.Anthropic()
    RESULTS_DIR.mkdir(parents=True, exist_ok=True)

    total = 0
    for dept, files in sorted(dept_files.items()):
        print(f"Processing {dept} ({len(files)} files for {program_name})")
        results = _analyze_files(client, files, RESULTS_DIR / f"{dept}.json")
        total += len(results)
        print(f"  {len(results)} total results for {dept}")

    print(f"\nDone. {total} total results across {len(dept_files)} departments for '{program_name}'")


def analyze_all(departments: list[str] | None = None):
    """Analyze syllabi for all (or specified) departments and save results."""
    client = anthropic.Anthropic()

    if departments is None:
        departments = sorted(
            {
                _extract_department(f.name)
                for f in SYLLABI_DIR.iterdir()
                if f.suffix.lower() in SUPPORTED_EXTENSIONS
                and _extract_department(f.name)
            }
        )

    if not departments:
        print("No syllabi files found in syllabi/")
        return

    for dept in departments:
        print(f"Processing department: {dept}")
        results = analyze_department(client, dept)
        print(f"  {len(results)} total results for {dept}")


if __name__ == "__main__":
    depts = sys.argv[1:] if len(sys.argv) > 1 else None
    analyze_all(depts)
